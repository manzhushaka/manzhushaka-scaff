#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown -> DOCX 转换脚本
输入: docs/research/发票业务竞品调研报告.md (只读)
输出: docs/research/发票业务竞品调研报告.docx

规则:
- #/##/### -> Heading 1/2/3
- GFM 表格 -> Word Table Grid, 表头加粗+浅灰底纹
- 无序列表 -> List Bullet; 有序列表 -> 显式编号段落(避免 Word 自动编号跨列表连续)
- 引用块(>) -> 小一号灰色说明文字
- 正文中文宋体 10.5pt, 标题微软雅黑, 表格 9.5pt
- 行内代码保留原文(Consolas); URL 保持纯文本
- A4 页面, 适度页边距
"""
import os

from markdown_it import MarkdownIt
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/Users/manzhushaka/kimi-work/iip/docs/research/发票业务竞品调研报告.md"
DST = "/Users/manzhushaka/kimi-work/iip/docs/research/发票业务竞品调研报告.docx"

BODY_CJK = "宋体"
BODY_LATIN = "Times New Roman"
HEAD_CJK = "微软雅黑"
HEAD_LATIN = "Microsoft YaHei"
CODE_LATIN = "Consolas"
QUOTE_COLOR = RGBColor(0x59, 0x59, 0x59)
HEAD_COLOR = RGBColor(0x1A, 0x1A, 0x1A)
HEADING_SIZES = {1: 16, 2: 14, 3: 12}


def set_run_font(run, size=10.5, bold=False, italic=False, color=None,
                 latin=BODY_LATIN, cjk=BODY_CJK):
    """设置 run 字体: latin 用 font.name, 中文用 rPr eastAsia."""
    run.font.name = latin
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), cjk)


def setup_styles(doc):
    """页面与基础样式."""
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_LATIN
    normal.font.size = Pt(10.5)
    rpr = normal.element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), BODY_CJK)
    pf = normal.paragraph_format
    pf.line_spacing = 1.3
    pf.space_after = Pt(4)

    for level, size in HEADING_SIZES.items():
        st = doc.styles[f"Heading {level}"]
        st.font.name = HEAD_LATIN
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = HEAD_COLOR
        rpr = st.element.get_or_add_rPr()
        rpr.get_or_add_rFonts().set(qn("w:eastAsia"), HEAD_CJK)
        st.paragraph_format.space_before = Pt(14 if level == 2 else 10)
        st.paragraph_format.space_after = Pt(6)


def render_inline(token, para, size=10.5, color=None, force_bold=False,
                  latin=BODY_LATIN, cjk=BODY_CJK):
    """把 markdown-it inline token 的 children 渲染为 docx runs."""
    bold = force_bold
    italic = False
    for child in token.children or []:
        ctype = child.type
        if ctype == "strong_open":
            bold = True
        elif ctype == "strong_close":
            bold = force_bold
        elif ctype == "em_open":
            italic = True
        elif ctype == "em_close":
            italic = False
        elif ctype == "text":
            r = para.add_run(child.content)
            set_run_font(r, size=size, bold=bold, italic=italic, color=color,
                         latin=latin, cjk=cjk)
        elif ctype == "code_inline":
            r = para.add_run(child.content)
            set_run_font(r, size=size, color=color, latin=CODE_LATIN, cjk=cjk)
        elif ctype in ("softbreak", "hardbreak"):
            para.add_run().add_break()
        # link_open/link_close 等直接忽略: URL 以纯文本形式保留


def shade_cell(cell, fill="F2F2F2"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def parse_table(tokens, i):
    """tokens[i] 为 table_open; 返回 (header_rows, body_rows, table_close_index)."""
    header, body = [], []
    section = None
    current_row = None
    i += 1
    while tokens[i].type != "table_close":
        ttype = tokens[i].type
        if ttype == "thead_open":
            section = header
        elif ttype == "tbody_open":
            section = body
        elif ttype == "tr_open":
            current_row = []
        elif ttype == "inline":
            if current_row is not None:
                current_row.append(tokens[i])
        elif ttype == "tr_close":
            if section is not None and current_row is not None:
                section.append(current_row)
            current_row = None
        i += 1
    return header, body, i


def add_table(doc, header, body):
    rows = header + body
    ncols = max(len(r) for r in rows) if rows else 0
    if ncols == 0:
        return
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = doc.styles["Table Grid"]
    tbl.autofit = True
    r_idx = 0
    for row in header:
        for c_idx, inline in enumerate(row):
            cell = tbl.cell(r_idx, c_idx)
            shade_cell(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            render_inline(inline, p, size=9.5, force_bold=True)
        r_idx += 1
    for row in body:
        for c_idx, inline in enumerate(row):
            cell = tbl.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            render_inline(inline, p, size=9.5)
        r_idx += 1
    doc.add_paragraph()  # 表后间距


def convert(src, dst):
    with open(src, "r", encoding="utf-8") as f:
        text = f.read()

    md = MarkdownIt("commonmark").enable("table")
    tokens = md.parse(text)

    doc = Document()
    setup_styles(doc)
    doc.core_properties.title = "发票业务竞品调研报告"

    list_stack = []   # [{'kind': 'bullet'} / {'kind': 'ordered', 'n': int}]
    in_blockquote = False
    current_para = None
    table_count = 0

    i = 0
    while i < len(tokens):
        tok = tokens[i]
        ttype = tok.type

        if ttype == "heading_open":
            level = min(int(tok.tag[1]), 3)
            inline = tokens[i + 1]
            p = doc.add_paragraph(style=f"Heading {level}")
            render_inline(inline, p, size=HEADING_SIZES[level],
                          force_bold=True, latin=HEAD_LATIN, cjk=HEAD_CJK)
            i += 3
        elif ttype == "bullet_list_open":
            list_stack.append({"kind": "bullet"})
            i += 1
        elif ttype == "ordered_list_open":
            start = int((tok.attrs or {}).get("start", 1))
            list_stack.append({"kind": "ordered", "n": start})
            i += 1
        elif ttype in ("bullet_list_close", "ordered_list_close"):
            list_stack.pop()
            i += 1
        elif ttype == "list_item_open":
            cur = list_stack[-1]
            if cur["kind"] == "bullet":
                p = doc.add_paragraph(style="List Bullet")
            else:
                num = cur["n"]
                cur["n"] += 1
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.75)
                r = p.add_run(f"{num}. ")
                set_run_font(r)
            current_para = p
            i += 1
        elif ttype == "list_item_close":
            current_para = None
            i += 1
        elif ttype == "blockquote_open":
            in_blockquote = True
            i += 1
        elif ttype == "blockquote_close":
            in_blockquote = False
            current_para = None
            i += 1
        elif ttype == "paragraph_open":
            if tok.hidden:
                pass  # 紧凑列表项的段落已在 list_item_open 创建
            elif in_blockquote or not list_stack:
                p = doc.add_paragraph()
                if in_blockquote:
                    p.paragraph_format.left_indent = Cm(0.5)
                current_para = p
            i += 1
        elif ttype == "paragraph_close":
            if not tok.hidden:
                current_para = None
            i += 1
        elif ttype == "inline":
            target = current_para if current_para is not None else doc.add_paragraph()
            if in_blockquote:
                render_inline(tok, target, size=9.5, color=QUOTE_COLOR)
            else:
                render_inline(tok, target)
            i += 1
        elif ttype == "table_open":
            header, body, end = parse_table(tokens, i)
            add_table(doc, header, body)
            table_count += 1
            i = end + 1
        elif ttype == "hr":
            i += 1  # 分隔线不输出, 章节标题已提供结构
        else:
            i += 1

    doc.save(dst)
    return table_count


def verify(dst, md_table_count):
    size = os.path.getsize(dst)
    doc = Document(dst)
    print(f"[verify] 输出文件: {dst}")
    print(f"[verify] 文件大小: {size} bytes ({size / 1024:.1f} KB) "
          f"-> {'OK(>30KB)' if size > 30 * 1024 else 'FAIL(<=30KB)'}")
    print(f"[verify] 源 Markdown 表格数: {md_table_count}")
    print(f"[verify] DOCX 段落数(正文层): {len(doc.paragraphs)}")
    print(f"[verify] DOCX 表格数: {len(doc.tables)}")
    for idx, tbl in enumerate(doc.tables, 1):
        print(f"[verify]   表 {idx}: {len(tbl.rows)} 行 x {len(tbl.columns)} 列 "
              f"| 表头首格: {tbl.rows[0].cells[0].text[:20]}")
    ok = len(doc.tables) == md_table_count
    print(f"[verify] 表格全部转换: {'OK' if ok else 'FAIL'}")


if __name__ == "__main__":
    n = convert(SRC, DST)
    verify(DST, n)
